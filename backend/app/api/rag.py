from io import BytesIO

import httpx
from docx import Document
from fastapi import APIRouter, Depends, File, Form, UploadFile
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.core.security import get_current_user
from app.models import DocumentChunk, UploadedDocument, User
from app.schemas import DocumentUploadResponse, SearchRequest, UrlIngestRequest
from app.services.rag_service import clean_evidence_text, clean_html_text, rag_service

router = APIRouter(prefix="/rag", tags=["rag"])


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    collection: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    raw = await file.read()
    filename = file.filename or "document.txt"
    pages = extract_pages(filename, raw)
    chunk_count = 0
    document = UploadedDocument(user_id=user.id, filename=filename, collection=collection, chunk_count=0)
    db.add(document)
    db.commit()
    db.refresh(document)
    for page_number, text in pages:
        before = len(rag_service._documents)
        chunk_count += rag_service.ingest(filename, collection, clean_evidence_text(text), page_number=page_number)
        for item in rag_service._documents[before:]:
            db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_id=item["chunk_id"],
                    source=filename,
                    page_number=page_number,
                    text=item["text"],
                    metadata_json={"collection": collection},
                )
            )
    document.chunk_count = chunk_count
    db.commit()
    db.refresh(document)
    return document


@router.post("/search")
def search(payload: SearchRequest, user: User = Depends(get_current_user)):
    return {"results": [hit.__dict__ for hit in rag_service.search(payload.query, payload.collection, payload.limit)]}


@router.post("/ingest-url", response_model=DocumentUploadResponse)
async def ingest_url(payload: UrlIngestRequest, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
        response = await client.get(payload.url, headers={"User-Agent": "VentureMindAI/1.0"})
        response.raise_for_status()
    text = clean_html_text(response.text)
    filename = str(response.url)
    document = UploadedDocument(user_id=user.id, filename=filename, collection=payload.collection, chunk_count=0, metadata_json={"url": filename})
    db.add(document)
    db.commit()
    db.refresh(document)
    before = len(rag_service._documents)
    chunk_count = rag_service.ingest(filename, payload.collection, text)
    for item in rag_service._documents[before:]:
        db.add(
            DocumentChunk(
                document_id=document.id,
                chunk_id=item["chunk_id"],
                source=filename,
                page_number=None,
                text=item["text"],
                metadata_json={"collection": payload.collection, "url": filename},
            )
        )
    document.chunk_count = chunk_count
    db.commit()
    db.refresh(document)
    return document


def extract_pages(filename: str, raw: bytes) -> list[tuple[int | None, str]]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        reader = PdfReader(BytesIO(raw))
        return [(index + 1, page.extract_text() or "") for index, page in enumerate(reader.pages)]
    if lower.endswith(".docx"):
        document = Document(BytesIO(raw))
        return [(None, "\n".join(paragraph.text for paragraph in document.paragraphs))]
    return [(None, raw.decode("utf-8", errors="ignore"))]


def strip_html(html: str) -> str:
    return clean_html_text(html)
