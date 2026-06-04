from io import BytesIO
import re
from textwrap import wrap
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from app.services.rag_service import clean_evidence_text, is_boilerplate_text


NOT_VERIFIED = "Not publicly verified from collected sources."

REPORT_SECTION_ORDER = [
    ("company_idea_overview", "Company / Idea Overview"),
    ("market_analysis", "Market Analysis"),
    ("competitor_analysis", "Competitor Analysis"),
    ("product_comparison", "Product Comparison"),
    ("pricing_comparison", "Pricing Comparison"),
    ("swot_analysis", "SWOT Analysis"),
    ("risks", "Risks"),
    ("strategic_recommendations", "Strategic Recommendations"),
    ("sources_used", "Sources Used"),
]

COMPANY_COMPARISON_SECTION_ORDER = [
    ("company_idea_overview", "Company / Idea Overview"),
    ("market_analysis", "Market Analysis"),
    ("competitor_analysis", "Competitor Analysis"),
    ("product_comparison", "Product Comparison"),
    ("pricing_comparison", "Pricing Comparison"),
    ("sources_used", "Sources Used"),
]


def format_business_report(raw_agent_output: dict[str, Any]) -> dict[str, Any]:
    if raw_agent_output.get("formatted_version") == 6:
        return raw_agent_output

    query = raw_agent_output.get("query") or raw_agent_output.get("input_text") or "this business question"
    sources = _unique_sources(raw_agent_output, query)
    mode = raw_agent_output.get("mode", "")
    entities = raw_agent_output.get("entities") or []
    confidence = _confidence(raw_agent_output, sources)
    market = raw_agent_output.get("market_intelligence", {})

    sections = {
        "company_idea_overview": _overview_section(query, mode, entities, sources, market),
        "market_analysis": _market_section(query, sources, market),
        "competitor_analysis": _competitor_section(query, mode, entities, sources, market),
        "product_comparison": _product_section(query, mode, entities, raw_agent_output, sources, market),
        "pricing_comparison": _pricing_section(sources),
        "sources_used": {
            "paragraphs": ["The report uses only the source links below. Numeric values are omitted when not directly verified."],
            "sources": [{"title": source["title"], "url": source["url"]} for source in sources],
        },
    }
    if mode != "company_comparison":
        sections.update(
            {
                "swot_analysis": _swot_section(raw_agent_output, query, sources),
                "risks": _list_section(_clean_items(raw_agent_output.get("sections", {}).get("risks", {})), "Risks could not be verified beyond collected evidence."),
                "strategic_recommendations": _strategy_section(raw_agent_output, query, sources, market),
            }
        )

    return {
        "formatted_version": 6,
        "title": raw_agent_output.get("title") or f"Business Strategy Report: {str(query)[:80]}",
        "query": query,
        "mode": mode,
        "entities": entities,
        "executive_summary": _executive_summary(query, mode, entities, sources, market),
        "sections": sections,
        "scorecard": {} if mode == "company_comparison" else _format_scorecard(raw_agent_output, sources, market),
        "viability_score": 0 if mode == "company_comparison" else _formatted_viability(raw_agent_output, sources, market),
        "confidence_score": confidence,
        "research_summary": raw_agent_output.get("research_summary", {}),
        "market_intelligence": market,
        "evidence_panel": {
            "sources": sources,
            "retrieved_chunks": _retrieved_chunks(raw_agent_output, query),
            "confidence": confidence,
        },
        "agent_panel": {
            "statuses": _agent_statuses(raw_agent_output),
        },
        "critic": {
            "confidence_score": confidence,
            "unsupported_claims": raw_agent_output.get("critic", {}).get("unsupported_claims", []),
        },
    }


def flatten_report(content: dict[str, Any]) -> list[tuple[str, str]]:
    formatted = format_business_report(content)
    rows = [("Executive Summary", formatted.get("executive_summary", ""))]
    section_order = COMPANY_COMPARISON_SECTION_ORDER if formatted.get("mode") == "company_comparison" else REPORT_SECTION_ORDER
    for key, title in section_order:
        rows.append((title, _section_to_text(formatted.get("sections", {}).get(key, {}))))
    if formatted.get("mode") != "company_comparison" and formatted.get("scorecard"):
        rows.append(("Business Scorecard", _scorecard_to_text(formatted.get("scorecard", {}))))
    return rows


def build_pdf(content: dict[str, Any]) -> bytes:
    rows = flatten_report(content)
    formatted = format_business_report(content)
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    y = height - 48
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(48, y, formatted.get("title", "VentureMind AI Report")[:90])
    y -= 32
    for title, body in rows:
        y = _pdf_heading(pdf, title, y, height)
        for paragraph in body.splitlines():
            if not paragraph.strip():
                y -= 8
                continue
            for line in wrap(paragraph, width=102):
                if y < 70:
                    pdf.showPage()
                    y = height - 48
                    pdf.setFont("Helvetica", 9)
                pdf.drawString(48, y, line)
                y -= 12
        y -= 10
    pdf.save()
    return buffer.getvalue()


def build_docx(content: dict[str, Any]) -> bytes:
    formatted = format_business_report(content)
    document = Document()
    document.add_heading(formatted.get("title", "VentureMind AI Report"), 0)
    for title, body in flatten_report(formatted):
        document.add_heading(title, level=1)
        for paragraph in body.splitlines():
            if paragraph.startswith("- "):
                document.add_paragraph(paragraph[2:], style="List Bullet")
            else:
                document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _executive_summary(query: str, mode: str, entities: list[str], sources: list[dict[str, Any]], market: dict[str, Any]) -> str:
    subject = _subject(query, entities)
    if not sources:
        return f"{subject} cannot be evaluated confidently because no verified sources were collected. The report avoids numeric claims until evidence is available."
    if mode == "company_comparison" and len(entities) >= 2:
        return (
            f"{subject} was compared using collected public sources. The available evidence supports a directional comparison of market presence, product positioning, pricing visibility, "
            f"and business signals, but exact financial metrics are shown only when directly verified."
        )
    competitors = [item["name"] for item in market.get("existing_competitors", [])]
    if competitors:
        return (
            f"{subject} is not a completely new market category. Collected sources show existing products such as {', '.join(competitors[:5])}. "
            f"The opportunity is therefore competitive rather than empty-market: the startup must differentiate by niche, workflow quality, pricing, integrations, or customer segment. "
            f"Financial figures remain unverified unless directly cited by sources."
        )
    return (
        f"{subject} shows directional potential based on collected public evidence. The strongest supported themes are market demand, product differentiation, and go-to-market focus; "
        f"financial figures remain unverified unless explicitly cited by sources."
    )


def _overview_section(query: str, mode: str, entities: list[str], sources: list[dict[str, Any]], market: dict[str, Any]) -> dict[str, Any]:
    subject = _subject(query, entities)
    if mode == "company_comparison" and len(entities) >= 2:
        paragraph = f"This report compares {subject} using public evidence from official, product, pricing, market, and business sources where available."
    elif market.get("existing_competitors"):
        names = ", ".join(item["name"] for item in market["existing_competitors"][:6])
        paragraph = f"{subject} appears to exist as an active market category. Collected sources identify existing products or competitors including {names}."
    else:
        paragraph = f"This report treats {subject} as a startup idea and checks public sources for similar products. No sufficiently verified competitor list was found from collected sources."
    return {"paragraphs": [paragraph, _source_sentence(sources)]}


def _market_section(query: str, sources: list[dict[str, Any]], market: dict[str, Any]) -> dict[str, Any]:
    bullets = _source_bullets(sources, ["market", "news", "business", "web"], limit=4)
    paragraphs = [
        "Market analysis is based on verified source snippets collected during web research. Exact market size, CAGR, and growth rates are included only when a collected source directly states them."
    ]
    category_terms = market.get("category_terms", [])
    if category_terms:
        paragraphs.append(f"The research agent searched this as: {', '.join(category_terms)}.")
    if not bullets:
        paragraphs.append(NOT_VERIFIED)
    return {"paragraphs": paragraphs, "bullets": bullets}


def _competitor_section(query: str, mode: str, entities: list[str], sources: list[dict[str, Any]], market: dict[str, Any]) -> dict[str, Any]:
    if mode == "company_comparison" and len(entities) >= 2:
        return {
            "paragraphs": [f"The comparison focuses on {', '.join(entities)} using public evidence. Missing fields are intentionally marked as not publicly verified."],
            "table": _side_by_side_table(entities, sources),
        }
    competitors = market.get("existing_competitors", [])
    if competitors:
        return {
            "paragraphs": [
                f"Existing products were found for {query}. This means the idea should be evaluated as a differentiated entrant, not as an untouched market."
            ],
            "table": [
                {
                    "competitor": item["name"],
                    "evidence": f"{item['name']} appears in collected market or product evidence for this category.",
                    "citation": {"label": item.get("source_title", "Source"), "url": item.get("source_url", "")},
                }
                for item in competitors[:8]
            ],
        }
    return {
        "paragraphs": [
            f"For {query}, the report uses industry-level competitive analysis. It does not create fake company metrics for a startup idea."
        ],
        "bullets": [
            "Differentiate through a focused niche, pricing model, workflow depth, geography, or underserved customer segment.",
            "Validate competitors with official product and pricing pages before final positioning decisions.",
        ],
    }


def _product_section(query: str, mode: str, entities: list[str], raw: dict[str, Any], sources: list[dict[str, Any]], market: dict[str, Any]) -> dict[str, Any]:
    product_sources = _filter_sources(sources, ["product", "official", "about"])
    if mode == "company_comparison" and len(entities) >= 2:
        return {
            "paragraphs": ["Product comparison uses collected official/product evidence. Unverified capabilities are not treated as confirmed."],
            "table": _side_by_side_table(entities, product_sources or sources, rows=("Positioning", "Product evidence", "Customer focus")),
        }
    competitors = market.get("existing_competitors", [])
    if competitors:
        return {
            "paragraphs": ["Similar products already exist, so the product plan should focus on a specific wedge rather than broad generic features."],
            "bullets": [
                f"Benchmark against {item['name']} using the cited source {item.get('source_title', 'Source')}."
                for item in competitors[:5]
            ],
        }
    return {
        "paragraphs": ["For a startup idea, product comparison is framed as capability planning rather than fake competitor metrics."],
        "bullets": _source_bullets(product_sources or sources, ["product", "official", "web"], limit=4)
        or ["Start with a narrow workflow, prove adoption, and expand only after customer validation."],
    }


def _pricing_section(sources: list[dict[str, Any]]) -> dict[str, Any]:
    pricing_sources = _filter_sources(sources, ["pricing"])
    if not pricing_sources:
        return {"paragraphs": [NOT_VERIFIED], "bullets": ["Collect official pricing pages before setting price benchmarks."]}
    return {
        "paragraphs": ["Pricing analysis uses official or pricing-related sources collected during research."],
        "bullets": _source_bullets(pricing_sources, ["pricing"], limit=5),
    }


def _financial_section(raw: dict[str, Any], sources: list[dict[str, Any]]) -> dict[str, Any]:
    raw_metrics = raw.get("sections", {}).get("financial_business_signals", {}).get("metrics", {})
    rows = []
    for metric in ["Revenue", "Revenue Growth", "CAGR", "Market Cap", "EBITDA", "Profit Margin", "Funding Raised", "Burn Rate", "Customer Growth"]:
        item = raw_metrics.get(metric, {})
        value = item.get("value")
        citation = item.get("citation", "")
        source_title = _source_title_for_url(sources, citation)
        if not value or str(value).lower() in {"not publicly verified", "none", ""} or not citation or not source_title:
            value = NOT_VERIFIED
        rows.append({"metric": metric, "value": value, "source": source_title or "No direct source"})
    return {
        "paragraphs": ["Financial and business signals are conservative. The report does not assign random revenue, CAGR, market cap, funding, EBITDA, or margin values."],
        "table": rows,
    }


def _swot_section(raw: dict[str, Any], query: str, sources: list[dict[str, Any]]) -> dict[str, Any]:
    swot = raw.get("sections", {}).get("swot_comparison", {})
    fallback = {
        "strengths": ["Clear opportunity can be evaluated through collected evidence."],
        "weaknesses": ["Some claims remain unverified from public sources."],
        "opportunities": ["Focus on a narrow customer segment and validate demand."],
        "threats": ["Incumbents can respond quickly if the market is attractive."],
    }
    return {
        key: _clean_items(swot.get(key, [])) or items
        for key, items in fallback.items()
    }


def _strategy_section(raw: dict[str, Any], query: str, sources: list[dict[str, Any]], market: dict[str, Any]) -> dict[str, Any]:
    competitors = [item["name"] for item in market.get("existing_competitors", [])]
    if raw.get("mode") == "company_comparison":
        entities = raw.get("entities") or []
        subject = " and ".join(entities) if len(entities) >= 2 else query
        items = [
            f"Use this as a directional comparison of {subject}; do not treat missing metrics as verified facts.",
            "Collect official product, pricing, investor, and annual-report sources for each company before making investment or strategy decisions.",
            "Compare each company on market position, product breadth, customer focus, delivery model, pricing visibility, and verified business signals.",
            "Where a source discusses only one company, keep that evidence in that company column and mark the other company as not publicly verified.",
        ]
    elif competitors:
        items = [
            f"Position the product against existing tools such as {', '.join(competitors[:4])} instead of presenting it as a new category.",
            _category_wedge_recommendation(query),
            _category_differentiation_recommendation(query),
            "Collect official pricing pages for the named competitors before deciding whether to compete on affordability or specialization.",
            _category_customer_discovery_recommendation(query),
        ]
    else:
        items = [
            "Start with a focused niche and validate demand before expanding product scope.",
            "Use official competitor and pricing evidence to decide positioning.",
            "Avoid financial planning based on unverified public numbers.",
        ]
    return {"paragraphs": ["Recommendations are based on collected evidence and explicitly avoid unsupported numeric assumptions."], "bullets": items[:6]}


def _category_wedge_recommendation(query: str) -> str:
    lowered = query.lower()
    if "hostel" in lowered:
        return "Choose a narrow wedge such as hostel-specific operations, student hostel administration, PG accommodation workflows, or budget-property management."
    if any(term in lowered for term in ("zepto", "blinkit", "quick commerce", "grocery")):
        return "Choose a narrow wedge such as a geography, delivery category, private-label assortment, subscription model, or underserved customer segment."
    return "Choose a narrow wedge by customer segment, geography, workflow depth, or pricing model."


def _category_differentiation_recommendation(query: str) -> str:
    lowered = query.lower()
    if "hostel" in lowered:
        return "Differentiate on workflows that generic hotel PMS tools may not handle deeply, such as bed allocation, mess/meal billing, visitor logs, fee reminders, and student/tenant records."
    if any(term in lowered for term in ("zepto", "blinkit", "quick commerce", "grocery")):
        return "Differentiate on delivery reliability, unit economics, assortment depth, local merchant partnerships, or customer retention rather than copying broad quick-commerce features."
    return "Differentiate on a specific workflow pain point that existing products do not solve well."


def _category_customer_discovery_recommendation(query: str) -> str:
    lowered = query.lower()
    if "hostel" in lowered:
        return "Run customer discovery with hostel owners, wardens, PG operators, and small hotel managers before building enterprise PMS features."
    if any(term in lowered for term in ("zepto", "blinkit", "quick commerce", "grocery")):
        return "Run customer discovery with target neighborhoods, frequent grocery buyers, delivery partners, and local merchants before making expansion assumptions."
    return "Run customer discovery with the specific buyer and user segments before expanding feature scope."


def _format_scorecard(raw: dict[str, Any], sources: list[dict[str, Any]], market: dict[str, Any]) -> dict[str, Any]:
    raw_scorecard = raw.get("scorecard", {})
    source_titles = [source["title"] for source in sources[:3]] or ["No cited evidence available"]
    theme = _evidence_theme(sources)
    competitor_count = len(market.get("existing_competitors", []))
    evidence_count = len(sources)
    labels = {
        "market_potential": "Market Potential",
        "innovation": "Innovation",
        "revenue_potential": "Revenue Potential",
        "competition_risk": "Competition Risk",
        "feasibility": "Feasibility",
        "scalability": "Scalability",
        "investment_attractiveness": "Investment Attractiveness",
    }
    formatted = {}
    for key, title in labels.items():
        score = _dimension_score(key, evidence_count, competitor_count)
        confidence = _average_confidence(sources)
        formatted[key] = {
            "title": title,
            "score": int(round(score)),
            "explanation": _score_explanation(title, theme, key, competitor_count),
            "evidence": source_titles,
            "confidence": confidence,
        }
    return formatted


def _formatted_viability(raw: dict[str, Any], sources: list[dict[str, Any]], market: dict[str, Any]) -> int:
    scorecard = _format_scorecard(raw, sources, market)
    return round(sum(item["score"] for item in scorecard.values()) / max(1, len(scorecard)))


def _dimension_score(key: str, evidence_count: int, competitor_count: int) -> int:
    if evidence_count == 0:
        if key == "competition_risk":
            return 50
        return 25
    evidence_bonus = min(18, evidence_count * 3)
    competition_pressure = min(28, competitor_count * 5)
    if key == "market_potential":
        return min(84, 52 + evidence_bonus + min(10, competitor_count * 2))
    if key == "innovation":
        return max(42, 72 + evidence_bonus // 2 - competition_pressure // 2)
    if key == "revenue_potential":
        return min(72, 45 + evidence_bonus + min(8, competitor_count))
    if key == "competition_risk":
        return min(88, 35 + competition_pressure + evidence_bonus // 2)
    if key == "feasibility":
        return min(78, 48 + evidence_bonus + min(8, competitor_count))
    if key == "scalability":
        return min(76, 46 + evidence_bonus + min(10, competitor_count))
    return min(74, 42 + evidence_bonus + min(10, competitor_count))


def _score_explanation(title: str, theme: str, key: str, competitor_count: int) -> str:
    if key == "competition_risk":
        return f"{title} reflects visible market overlap and incumbent pressure from {competitor_count} detected competitor signals. {theme}"
    if key == "revenue_potential":
        return f"{title} reflects demand signals only; exact revenue is not assumed without direct verification. {theme}"
    if key == "innovation" and competitor_count:
        return f"{title} is moderated because similar products already exist, so differentiation matters more than novelty. {theme}"
    return f"{title} reflects the strength, relevance, and confidence of collected evidence. {theme}"


def _unique_sources(raw: dict[str, Any], query: str) -> list[dict[str, Any]]:
    candidates = []
    candidates.extend(raw.get("web_sources", []))
    candidates.extend(raw.get("citations", []))
    candidates.extend(raw.get("evidence_panel", {}).get("sources", []))
    seen = set()
    sources = []
    for item in candidates:
        url = (item.get("url") or item.get("source_url") or "").strip()
        title = _clean_text(item.get("title") or item.get("source") or item.get("source_name") or "Source")
        snippet = _clean_text(item.get("snippet") or item.get("text") or item.get("claim") or "")
        if not url and not snippet:
            continue
        if not _source_relevant_to_query(query, title, snippet, url):
            continue
        key = url or f"{title}:{snippet[:80]}"
        if key in seen or _incomplete(snippet):
            continue
        seen.add(key)
        sources.append(
            {
                "title": title[:180],
                "url": url,
                "snippet": snippet[:500],
                "source_type": item.get("source_type") or item.get("collection") or "web",
                "confidence": float(item.get("confidence_score") or item.get("confidence") or 0.5),
                "collected_at": item.get("collected_at") or item.get("timestamp") or "",
            }
        )
    return sources[:20]


def _source_relevant_to_query(query: str, title: str, snippet: str, url: str) -> bool:
    anchors = set(_tokens(query)) - {
        "app",
        "business",
        "company",
        "compare",
        "comparison",
        "idea",
        "management",
        "market",
        "platform",
        "product",
        "software",
        "startup",
        "system",
        "the",
        "vs",
    }
    if not anchors:
        return True
    text = f"{title} {snippet} {url}".lower()
    if any(name in query.lower() for name in ("zepto", "blinkit")) and not any(
        marker in text for marker in ("zeptonow", "blinkit", "zomato", "quick commerce", "grocery", "delivery")
    ):
        return False
    if "hostel" in query.lower() and not any(marker in text for marker in ("hostel", "hotel", "property management", "pms", "cloudbeds", "webrezpro", "sirvoy", "beds24")):
        return False
    haystack = set(_tokens(text))
    return bool(anchors & haystack)


def _retrieved_chunks(raw: dict[str, Any], query: str) -> list[dict[str, Any]]:
    chunks = []
    for item in raw.get("citations", [])[:10]:
        if item.get("collection") == "Web Research":
            continue
        text = _clean_text(item.get("text", ""))
        if text and not is_boilerplate_text(text) and not _incomplete(text) and _source_relevant_to_query(query, item.get("source", ""), text, ""):
            chunks.append({"source": item.get("source", "Uploaded evidence"), "text": text[:350], "confidence": item.get("confidence", 0.5)})
    return chunks


def _agent_statuses(raw: dict[str, Any]) -> list[dict[str, str]]:
    rows = raw.get("agent_statuses") or raw.get("agent_panel", {}).get("statuses") or []
    aliases = {
        "web_research": "Research Agent",
        "research": "Research Agent",
        "competitor": "Competitor Agent",
        "swot": "SWOT Agent",
        "critic": "Critic Agent",
        "report": "Report Agent",
    }
    ordered = ["web_research", "competitor", "swot", "critic", "report"]
    by_agent = {row.get("agent"): row.get("status", "pending") for row in rows}
    result = []
    for key in ordered:
        result.append({"agent": aliases[key], "status": _status_label(by_agent.get(key, "completed" if rows else "pending"))})
    return result


def _status_label(value: str) -> str:
    return str(value or "pending").replace("_", " ").title()


def _source_sentence(sources: list[dict[str, Any]]) -> str:
    if not sources:
        return "No verified public source was collected for this section."
    names = ", ".join(source["title"] for source in sources[:3])
    return f"Key collected sources include {names}."


def _source_bullets(sources: list[dict[str, Any]], source_types: list[str], limit: int = 4) -> list[str]:
    bullets = []
    for source in _filter_sources(sources, source_types):
        bullet = _snippet_to_bullet(source)
        if bullet:
            bullets.append(bullet)
        if len(bullets) >= limit:
            break
    return bullets


def _filter_sources(sources: list[dict[str, Any]], source_types: list[str]) -> list[dict[str, Any]]:
    needles = [item.lower() for item in source_types]
    return [source for source in sources if any(needle in f"{source.get('source_type', '')} {source.get('title', '')} {source.get('snippet', '')}".lower() for needle in needles)]


def _snippet_to_bullet(source: dict[str, Any]) -> str:
    snippet = _clean_text(source.get("snippet", ""))
    if not snippet or is_boilerplate_text(snippet):
        return ""
    useful_sentences = [sentence for sentence in _sentences(snippet) if not is_boilerplate_text(sentence)]
    sentence = useful_sentences[0].strip() if useful_sentences else ""
    if len(sentence) < 30:
        sentence = snippet[:180].strip()
    if not sentence or is_boilerplate_text(sentence) or not _meaningful_evidence_sentence(sentence):
        return ""
    return f"{sentence.rstrip('.')} ({source.get('title', 'Source')})."


def _side_by_side_table(entities: list[str], sources: list[dict[str, Any]], rows: tuple[str, ...] = ("Market evidence", "Product evidence", "Pricing evidence")) -> list[dict[str, str]]:
    table = []
    for row_name in rows:
        row = {"item": row_name}
        for entity in entities:
            row[entity] = _entity_evidence_text(entity, row_name, sources) or NOT_VERIFIED
        table.append(row)
    return table


def _entity_evidence_text(entity: str, row_name: str, sources: list[dict[str, Any]]) -> str:
    source = _best_entity_source(entity, row_name, sources)
    if not source:
        return ""
    sentence = source.get("sentence") or _clean_sentence(source.get("snippet", ""))
    if not sentence:
        return ""
    return f"{sentence} ({source.get('title', 'Source')})."


def _best_entity_source(entity: str, row_name: str, sources: list[dict[str, Any]]) -> dict[str, Any] | None:
    entity_l = entity.lower()
    terms_by_row = {
        "market evidence": ["market", "share", "growth", "valuation", "investor", "position", "dominant"],
        "product evidence": ["product", "platform", "delivery", "service", "vertical", "quick", "instamart", "blinkit"],
        "pricing evidence": ["pricing", "price", "fee", "commission", "subscription", "plan"],
        "positioning": ["position", "market", "focus", "brand", "customer"],
        "customer focus": ["customer", "user", "consumer", "merchant", "restaurant"],
    }
    terms = terms_by_row.get(row_name.lower(), row_name.lower().split())
    fallback = None
    for source in sources:
        for sentence in _sentences(source.get("snippet", "")):
            lowered = sentence.lower()
            if entity_l not in lowered:
                continue
            candidate = {**source, "sentence": _clean_sentence(sentence)}
            if any(term in lowered for term in terms):
                return candidate
            fallback = fallback or candidate
    return fallback if row_name.lower() in {"market evidence", "positioning", "customer focus"} else None


def _sentences(text: str) -> list[str]:
    cleaned = _clean_text(text)
    chunks = re.split(r"(?<=[.!?])\s+|\s+##+\s+|\s+###\s+", cleaned)
    return [chunk.strip(" -#") for chunk in chunks if len(chunk.strip()) > 20 and not is_boilerplate_text(chunk)]


def _clean_sentence(text: str) -> str:
    cleaned = _clean_text(text).replace("### ", "").replace("## ", "")
    cleaned = cleaned.replace("â€™", "'").replace("Â·", "-")
    return cleaned[:260].rstrip(" ,;:-")


def _meaningful_evidence_sentence(text: str) -> bool:
    lowered = text.lower()
    if len(lowered) < 55 and not any(term in lowered for term in ("market", "growth", "delivery", "pricing", "revenue", "customer", "competitor", "platform", "business")):
        return False
    if "reports, statistics" in lowered and "market" not in lowered:
        return False
    action_terms = ("is ", "are ", "has ", "have ", "shows", "reports", "states", "grew", "growth", "offers", "provides", "operates", "delivers", "launched")
    return any(term in lowered for term in action_terms)


def _clean_items(value: Any) -> list[str]:
    items: list[str] = []
    if isinstance(value, str):
        return [_clean_text(value)] if value.strip() else []
    if isinstance(value, list):
        for item in value:
            items.extend(_clean_items(item))
    if isinstance(value, dict):
        preferred = value.get("item") or value.get("recommendation") or value.get("summary") or value.get("reason") or value.get("gap") or value.get("action")
        if preferred:
            items.append(_clean_text(str(preferred)))
        else:
            for item in value.values():
                if isinstance(item, (list, dict)):
                    items.extend(_clean_items(item))
    deduped = []
    seen = set()
    for item in items:
        if len(item) < 20 or _incomplete(item):
            continue
        key = item.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def _list_section(items: list[str], fallback: str) -> dict[str, Any]:
    return {"bullets": items[:6] or [fallback]}


def _section_to_text(section: Any) -> str:
    if isinstance(section, str):
        return section
    if not isinstance(section, dict):
        return str(section or "")
    lines: list[str] = []
    for paragraph in section.get("paragraphs", []):
        lines.append(paragraph)
    for bullet in section.get("bullets", []):
        lines.append(f"- {bullet}")
    for key in ("strengths", "weaknesses", "opportunities", "threats"):
        if section.get(key):
            lines.append(f"{key.title()}:")
            lines.extend(f"- {item}" for item in section[key])
    for row in section.get("table", []):
        lines.append(" | ".join(f"{key}: {value}" for key, value in row.items()))
    for source in section.get("sources", []):
        lines.append(f"- {source.get('title', 'Source')}: {source.get('url', '')}")
    return "\n".join(lines)


def _scorecard_to_text(scorecard: dict[str, Any]) -> str:
    lines = []
    for item in scorecard.values():
        lines.append(f"{item.get('title', 'Score')}: {item.get('score', 0)}/100")
        lines.append(f"Explanation: {item.get('explanation', '')}")
        lines.append("Evidence:")
        lines.extend(f"- {source}" for source in item.get("evidence", []))
        lines.append(f"Confidence: {round(float(item.get('confidence', 0)) * 100)}%")
    return "\n".join(lines)


def _pdf_heading(pdf: canvas.Canvas, title: str, y: float, height: float) -> float:
    if y < 90:
        pdf.showPage()
        y = height - 48
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(48, y, title)
    pdf.setFont("Helvetica", 9)
    return y - 16


def _subject(query: str, entities: list[str]) -> str:
    return " vs ".join(entities) if entities else query


def _confidence(raw: dict[str, Any], sources: list[dict[str, Any]]) -> float:
    if not sources:
        return 0.25
    return float(raw.get("critic", {}).get("confidence_score") or raw.get("confidence_score") or _average_confidence(sources))


def _average_confidence(sources: list[dict[str, Any]]) -> float:
    if not sources:
        return 0.25
    return round(sum(float(source.get("confidence", 0.5)) for source in sources) / len(sources), 2)


def _evidence_theme(sources: list[dict[str, Any]]) -> str:
    if not sources:
        return "No verified evidence was collected, so confidence is low."
    types = {source.get("source_type", "web") for source in sources[:5]}
    return f"Evidence came from {', '.join(sorted(types))} sources with average confidence of {round(_average_confidence(sources) * 100)}%."


def _source_title_for_url(sources: list[dict[str, Any]], url: str) -> str:
    if not url:
        return ""
    for source in sources:
        if source.get("url") == url:
            return source.get("title", "")
    return ""


def _clean_text(value: str) -> str:
    return clean_evidence_text(str(value or "").replace("\n", " ").replace("\r", " "))


def _tokens(value: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", str(value or "").lower())


def _incomplete(text: str) -> bool:
    cleaned = _clean_text(text)
    if not cleaned:
        return True
    bad_markers = ["...", "read more", "enable javascript", "unsupported browser", "captcha"]
    if any(marker in cleaned.lower() for marker in bad_markers):
        return True
    return False
